from pathlib import Path
import shutil
import hashlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(
    r"C:\Users\Desarrollo\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-experiment-analysis\assets\reference.docx"
)
OUT = Path(r"C:\Users\Desarrollo\Downloads\Manual de Funcionalidades Aula Virtual - Actualizado.docx")
BUILD = Path(r"C:\xampp\htdocs\aula-virtual-docker\manual-build")

PRIMARY = "1F6AE1"
PRIMARY_SOFT = "EAF2FF"
HEADING = "0B253F"
MUTED = "5B6B80"
BORDER = "D9E2EC"
GREEN = "DCFCE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_runs(paragraph, parts):
    for text, opts in parts:
        run = paragraph.add_run(text)
        run.font.name = "Georgia"
        run.bold = bool(opts.get("bold"))
        if opts.get("color"):
            run.font.color.rgb = RGBColor.from_string(opts["color"])
        if opts.get("size"):
            run.font.size = Pt(opts["size"])


def main():
    BUILD.mkdir(parents=True, exist_ok=True)
    sha = hashlib.sha256(BASE.read_bytes()).hexdigest()
    (BUILD / "artifact.md").write_text(
        "\n".join(
            [
                "# Artifact Template Contract - Aula Virtual Manual",
                f"Reference: {BASE}",
                f"SHA-256: {sha}",
                "Page system: Letter 8.5 x 11 in, portrait, 1 in margins, one section.",
                "Typography evidence: reference uses Georgia as dominant font; generated manual reuses the copied DOCX and applies Georgia-based styles.",
                "Editable slot strategy: body content is replaced with the Aula Virtual manual while preserving the reference package as source, section page setup, styles, headers and footers.",
                f"Output: {OUT}",
            ]
        ),
        encoding="utf-8",
    )

    shutil.copyfile(BASE, OUT)
    doc = Document(str(OUT))

    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    for style_name, size in [
        ("Normal", 10.5),
        ("Title", 26),
        ("Subtitle", 13),
        ("Heading 1", 18),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Georgia"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(HEADING)
        except Exception:
            pass

    def para(text="", style=None, bold=False, color=None, size=None, align=None, space_after=6):
        p = doc.add_paragraph(style=style)
        if text:
            add_runs(p, [(text, {"bold": bold, "color": color or HEADING, "size": size})])
        p.paragraph_format.space_after = Pt(space_after)
        if align:
            p.alignment = align
        return p

    def heading(text, level=1):
        return para(text, style=f"Heading {level}", bold=True, color=HEADING, space_after=8)

    def bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3)
        add_runs(p, [("• ", {"color": PRIMARY}), (text, {"color": HEADING})])
        return p

    def callout(title, body_text, fill=PRIMARY_SOFT):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.cell(0, 0)
        set_cell_shading(cell, fill)
        set_cell_border(cell, "BFD7FF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        add_runs(p, [(title, {"bold": True, "color": PRIMARY, "size": 11})])
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        add_runs(p2, [(body_text, {"color": HEADING, "size": 10})])
        doc.add_paragraph()

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_shading(hdr[i], PRIMARY_SOFT)
            set_cell_border(hdr[i])
            hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            add_runs(hdr[i].paragraphs[0], [(h, {"bold": True, "color": HEADING, "size": 9})])
        for row in rows:
            cells = t.add_row().cells
            for i, value in enumerate(row):
                set_cell_border(cells[i])
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                add_runs(cells[i].paragraphs[0], [(str(value), {"color": HEADING, "size": 8.7})])
        doc.add_paragraph()

    para("Manual de Funcionalidades", style="Title", bold=True, color=HEADING, size=28, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para("Aula Virtual Smart Data", style="Subtitle", color=PRIMARY, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para("Estado actual, flujos operativos y plan de implementación", color=MUTED, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    table(
        ["Campo", "Detalle"],
        [
            ["Proyecto", "aula-virtual + aula-virtual-api-servicios"],
            ["Entorno analizado", r"C:\xampp\htdocs\aula-virtual-docker"],
            ["Fecha de actualización", "12 de agosto de 2026"],
            ["Audiencia", "Jefatura, producto, operaciones académicas, soporte y desarrollo"],
            ["Fuente", "Manual original, revisión del proyecto local y mejoras implementadas durante el ciclo de trabajo"],
        ],
    )

    callout(
        "Resumen ejecutivo",
        "El Aula Virtual ya cuenta con una base funcional amplia para alumnos, docentes y administradores: cursos, sesiones, materiales, videos, evaluaciones, encuestas, asistencia, calificaciones y certificados. La siguiente etapa debe enfocarse en estabilización productiva, automatización de procesos recurrentes, rendimiento, trazabilidad y cierre de flujos operativos.",
    )

    heading("1. Alcance Del Manual")
    para("Este documento actualiza el manual funcional del Aula Virtual y lo convierte en una guía de estado del proyecto. Su objetivo es explicar qué funciona hoy, qué módulos están en estabilización y qué mejoras deben priorizarse para producción.")
    bullet("Describe los flujos principales para alumno, docente y administrador.")
    bullet("Resume integraciones con API Servicios, Google Drive, Zoom, Gen-Docs y SGA.")
    bullet("Diferencia funcionalidades operativas, funcionalidades en consolidación y pendientes de segunda fase.")
    bullet("Propone un roadmap presentable para jefatura, con criterios de aceptación por módulo.")

    heading("2. Arquitectura Funcional")
    table(
        ["Componente", "Responsabilidad", "Estado"],
        [
            ["aula-virtual", "Portal Laravel para login, experiencia del alumno, backoffice, vistas, carga diferida y UX.", "Operativo con mejoras recientes de diseño, rendimiento y accesibilidad."],
            ["aula-virtual-api-servicios", "API Lumen que conecta el portal con la base académica, cursos, sesiones, evaluaciones, encuestas, certificados y asistencia.", "Operativo; requiere completar endpoints agregados y medición SQL en producción."],
            ["Base smartdata / mysql_cursos", "Datos académicos: cursos, sesiones, alumnos, materiales, evaluaciones, asistencia y certificados.", "Local alineado parcialmente; algunas migraciones dependen de conexión MySQL disponible."],
            ["Gen-Docs", "Fuente de verdad para formularios y respuestas de encuestas.", "Integración funcional planteada e iniciada; debe cerrar validación de duplicidad y resultados."],
            ["SGA Web", "Fuente de generación de diplomas/certificados.", "Aula consume diplomas; SGA conserva generación y diseño."],
            ["Google Drive", "Almacenamiento de videos, archivos de sesión, materiales y chat TXT de Zoom.", "Carga manual funcionando con ajustes; automatización Zoom queda como fase posterior."],
            ["Zoom", "Fuente de reuniones, accesos y futura conciliación de asistencia.", "Meetings integrado a nivel de modelo; webhooks/reportes requieren despliegue controlado."],
        ],
    )

    heading("3. Estado General Por Módulo")
    table(
        ["Módulo", "Ya Funciona", "Pendiente Recomendado"],
        [
            ["Autenticación y roles", "Login por usuario; navegación diferenciada para alumno, docente y administrador.", "Endurecer mensajes de error, auditoría de accesos y pruebas de permisos por rol."],
            ["Cursos y sesiones", "Listado, detalle, navegación por sesión, filtros, búsqueda, progreso y tabs.", "Unificar completamente experiencia alumno/docente/admin y medir tiempos por ruta."],
            ["Videos de sesión", "Subida manual a Drive por chunks; visualización de grabación y TXT de chat Zoom asociado.", "Automatizar ingesta desde Zoom Cloud Recordings y reducir fallos 413/chunk."],
            ["Materiales", "Subida, edición, vista previa de imagen/PDF, descarga autorizada y múltiples archivos por sesión.", "Auditar formatos, nombres correctos y limpieza de archivos huérfanos."],
            ["Anuncios y comunidad", "Anuncios por curso/sesión y comunidad con comentarios.", "Mejorar notificaciones, no leídos y moderación."],
            ["Evaluaciones y trabajos", "Exámenes, trabajos prácticos, rúbricas, entregas, borrador/finalización y calificación.", "Consolidar plan por hitos y reglas de fechas límite por sesión final."],
            ["Encuestas", "Integración con esquema Gen-Docs y formulario nativo para alumno.", "Cerrar estado respondida, resultados avanzados, privacidad docente y exportación."],
            ["Asistencia", "Modelo híbrido por rol diseñado y vistas base implementadas.", "Activar webhooks Zoom, conciliación, reportes y corrección manual en modo observación."],
            ["Calificaciones", "Vista docente/admin para revisar notas y calificar trabajos.", "Mejorar consistencia de estados alumno: entregado, calificado, nota publicada."],
            ["Certificados", "Integración con diplomas SGA y apartado Mi certificado para alumno.", "Sincronización robusta, envío por correo y trazabilidad final."],
            ["Docker y rendimiento", "Docker local con PHP-FPM/Nginx/Redis iniciado como objetivo.", "Terminar despliegue VPS, OPcache, workers, scheduler y métricas TTFB."],
        ],
    )

    heading("4. Roles Y Accesos")
    table(
        ["Rol", "Puede Ver", "Puede Ejecutar", "Restricciones"],
        [
            ["Alumno", "Sus cursos, sesiones, materiales publicados, evaluaciones asignadas, encuestas disponibles, asistencia propia y certificado propio.", "Responder encuestas, rendir evaluaciones, subir entregas, descargar materiales, participar en comunidad.", "No ve acciones administrativas, respuestas de otros alumnos, asistencia ajena ni datos internos."],
            ["Docente", "Cursos asignados, sesiones, materiales, evaluaciones, encuestas agregadas, calificaciones y asistencia de sus cursos.", "Subir contenido, configurar actividades, revisar entregas y corregir alumnos cuando corresponda.", "No debe corregir su propia asistencia ni ver datos sensibles no necesarios."],
            ["Administrador", "Todos los cursos, responsables, asistencia, certificados, encuestas, calificaciones y estados operativos.", "Supervisar, sincronizar, exportar, conciliar y corregir según módulo.", "Debe conservar trazabilidad y evitar exponer credenciales, correos o URLs sensibles en logs."],
        ],
    )

    heading("5. Flujo Del Alumno")
    heading("5.1 Pantalla Inicial Mis Cursos", 2)
    para("La pantalla inicial del alumno fue orientada a continuidad. Debe mostrar tres tarjetas por fila en escritorio, tabs de En progreso, Completados y Sugeridos, búsqueda por curso/docente/edición y un CTA principal por curso.")
    bullet("El progreso debe calcularse sobre sesiones reales del curso, por ejemplo 6 de 15 y no 12 de 30.")
    bullet("La pestaña Sugeridos debe proponer cursos relacionados con lo que el alumno lleva o llevó, sin inscribir ni iniciar pagos automáticamente.")
    bullet("Cada tarjeta debe responder: qué curso es, quién enseña, cuánto avanzó y cuál es el siguiente paso.")

    heading("5.2 Aula Del Curso", 2)
    para("El aula del alumno está alineada al patrón moderno del backoffice: sesiones a la izquierda, contenido central por pestañas y comunidad plegable. El alumno solo ve contenido publicado y autorizado.")
    table(
        ["Pestaña", "Uso Para Alumno", "Estado UX Esperado"],
        [
            ["Video", "Ver grabación de clase y chat TXT de Zoom.", "Grabación disponible, todavía no disponible o enlace externo seguro."],
            ["Materiales", "Consultar recursos publicados y descargar archivos.", "Recursos de la sesión o estado vacío claro."],
            ["Evaluaciones", "Ver solo actividades reales de la sesión.", "Oculta si la sesión no tiene examen/trabajo; muestra nota si ya está publicada."],
            ["Encuestas", "Responder encuestas de sesión o final.", "Responder, disponible pronto o respuesta enviada."],
            ["Anuncios", "Leer comunicados relevantes.", "No hay anuncios nuevos o listado de comunicados."],
            ["Mi asistencia", "Consultar su propio estado de asistencia.", "Pendiente de Zoom, asistió, no asistió, justificada o no aplica."],
        ],
    )

    heading("5.3 Evaluaciones Y Trabajos", 2)
    para("Las evaluaciones no deben asumirse para todas las sesiones. La experiencia correcta es por hitos: parcial, final, avances, proyecto integrador o presentación. En sesiones sin actividad evaluable la pestaña puede ocultarse para reducir ruido.")
    bullet("Trabajo parcial o examen parcial: plazo hasta el día de la última sesión.")
    bullet("Trabajo final o examen final: plazo sugerido de cinco días después de la última sesión.")
    bullet("El docente puede ampliar el plazo por excepción, manteniendo trazabilidad.")
    bullet("El alumno debe ver fecha límite, puntaje máximo, peso, estado de entrega y nota publicada cuando exista.")

    heading("6. Flujo Docente Y Administrador")
    heading("6.1 Gestión De Cursos", 2)
    para("La vista de cursos debe ser un panel operativo: métricas superiores, búsqueda por Enter, tabs por estado y tarjetas compactas. Para administrador se recomienda nombrar el módulo como Cursos o Gestión de cursos, no Mis Cursos, porque supervisa cursos de toda la operación.")

    heading("6.2 Contenido De Sesión", 2)
    table(
        ["Acción", "Estado Actual", "Mejora Recomendada"],
        [
            ["Subir video", "Carga manual con chunks hacia Google Drive.", "Automatizar lectura de Zoom Cloud Recordings por cuenta/curso y asociar MP4/TXT por patrón de nombre."],
            ["Chat de Zoom", "TXT asociado al video como Chat de la clase.", "Permitir vista legible por hora, participante y mensaje; no convertirlo en material."],
            ["Materiales", "Múltiples recursos por sesión, vista previa y descarga.", "Agrupar por tipo y mostrar acciones consistentes para alumno/docente."],
            ["Anuncios", "Anuncios de curso y sesión.", "Mejorar notificaciones y confirmaciones no intrusivas."],
        ],
    )

    heading("6.3 Evaluaciones, Calificación Y Notas", 2)
    para("El flujo actual permite crear evaluaciones, trabajos, rúbricas, publicar, recibir entregas y calificar. La mejora estratégica es convertirlo en un Plan de evaluación del curso, donde cada hito queda asignado a una sesión específica y no se muestran pendientes falsos.")
    bullet("Plantillas rápidas: parcial + final; proyecto integrador con avances; cuatro entregables por sesiones específicas.")
    bullet("La calificación debe actualizar inmediatamente la vista del alumno y resolver inconsistencias como Calificado pero Nota pendiente.")
    bullet("La pantalla de calificación debe mantener guardado robusto, siguiente entrega y estados de error recuperables.")

    heading("7. Encuestas")
    para("La estrategia correcta es usar Gen-Docs como fuente de verdad para formularios, preguntas, enlaces y respuestas. Aula Virtual debe mostrar el formulario nativo para el alumno y resultados agregados para docente/admin.")
    table(
        ["Flujo", "Decisión", "Próximo Cierre"],
        [
            ["Alumno responde", "Autenticado, sin pedir correo nuevamente.", "Corregir cache/estado para que después de responder aparezca Respuesta enviada."],
            ["Docente consulta", "Solo agregados y privacidad por muestra mínima.", "Resultados accionables por pregunta, sesión y docente."],
            ["Administrador consulta", "Puede ver auditoría e identidad para soporte.", "Exportación CSV con filtros y UTF-8."],
            ["Datos", "Tablas plurales de Gen-Docs.", "Retirar sistema paralelo singular vacío cuando ya no haya lecturas legacy."],
        ],
    )

    heading("8. Asistencia")
    para("El plan de asistencia es híbrido y diferenciado por rol: alumnos no tienen tardanza ni permanencia mínima; docentes requieren puntualidad y 80% de permanencia. El clic en Zoom registra intento, pero la asistencia válida debe confirmarse con Zoom.")
    table(
        ["Participante", "Regla Académica", "Estados"],
        [
            ["Alumno", "Asiste si Zoom confirma al menos un ingreso real; duración solo evidencia.", "Asistió, No asistió, Justificada, Pendiente, No aplica."],
            ["Docente", "Presente si entra hasta 7:00:59 y cumple 80%; desde 7:01:00 es tardanza si cumple 80%.", "Presente, Tardanza, Falta, Justificada, Pendiente, No aplica."],
            ["Administrador", "Supervisa, concilia, corrige e identifica participantes.", "Debe auditar todo cambio manual con motivo."],
        ],
    )

    heading("9. Certificados Y Diplomas")
    para("SGA debe seguir generando diplomas; Aula Virtual debe consultarlos, mostrarlos y operar envío/sincronización. Para alumno se recomienda mantener el certificado dentro del curso, como Mi certificado, porque es el contexto natural donde espera encontrarlo.")
    bullet("Alumno ve estado: no disponible, en preparación, disponible, enviado o requiere revisión.")
    bullet("Administrador gestiona por curso: sincronizar desde SGA, ver diploma, descargar, copiar enlace y enviar por correo.")
    bullet("No se deben exponer rutas físicas, correos de otros alumnos ni datos administrativos.")

    heading("10. Rendimiento, Docker Y Producción")
    para("El proyecto ya inició una transición hacia un entorno más cercano a producción con Docker, Redis, PHP-FPM, Nginx, workers y scheduler. Esta línea debe cerrarse antes del despliegue VPS para evitar diferencias entre local y producción.")
    table(
        ["Capa", "Acción Recomendada", "Objetivo"],
        [
            ["Runtime", "APP_DEBUG=false, OPcache, config/route/view cache, composer optimize autoloader.", "Reducir latencia base y errores visibles."],
            ["Cache", "Redis para sesiones, cache de cursos, paneles y respuestas API.", "Vistas cacheadas por debajo de 500 ms."],
            ["API", "Endpoints agregados por alumno/backoffice/workspace.", "Menos llamadas repetidas y menos N+1."],
            ["SQL", "Medir queries lentas y agregar índices finales por evidencia.", "TTFB frío menor a 1.5-2 s."],
            ["Frontend", "Carga diferida por pestaña, skeleton local y sessionStorage.", "Cambios de pestaña/sesión menores a 300 ms si ya fueron visitados."],
            ["Docker/VPS", "Nginx + PHP-FPM + Redis + workers + scheduler + healthchecks.", "Ambiente reproducible y operable."],
        ],
    )

    heading("11. Automatización Recomendable: Videos Desde Zoom")
    para("La automatización más valiosa para operación académica es importar grabaciones desde Zoom. Cada curso-edición usa cuentas de Zoom registradas en meetings y los archivos siguen un patrón como Sesión 14-G1-Arquitectura de Datos End-to-End.")
    bullet("Proceso sugerido: job programado consulta Zoom Cloud Recordings por cuenta, filtra MP4 y TXT, asocia por sesión/grupo/curso y registra en Aula.")
    bullet("Solo se deben subir MP4 y TXT; el TXT queda como Chat de la clase asociado al video.")
    bullet("La asociación ambigua no debe publicar automáticamente; debe quedar en revisión administrativa.")
    bullet("El sistema debe registrar estado: detectado, asociado, publicado, fallido o requiere revisión.")

    heading("12. Roadmap Propuesto Para Jefatura")
    table(
        ["Fase", "Tiempo Estimado", "Objetivo", "Entregable"],
        [
            ["Fase 1 - Estabilización", "1 a 2 semanas", "Resolver errores 500, estados inconsistentes, cache de encuestas/evaluaciones y descargas.", "Portal estable local/Docker con pruebas críticas."],
            ["Fase 2 - UX y manual operativo", "1 semana", "Pulir textos no técnicos, estados vacíos, navegación alumno/docente/admin y documentación.", "Manual actualizado, checklist QA y pantallas consistentes."],
            ["Fase 3 - Producción VPS", "1 a 2 semanas", "Deploy con Nginx, PHP-FPM, Redis, workers, scheduler, backups y logs.", "Ambiente productivo con healthchecks y rollback."],
            ["Fase 4 - Automatización Zoom", "2 a 3 semanas", "Importar MP4/TXT automáticamente desde cuentas Zoom por curso-edición.", "Job automático, pantalla de revisión y publicación segura."],
            ["Fase 5 - Analítica y reportes", "2 semanas", "Resultados avanzados de encuestas, asistencia conciliada y certificados trazables.", "Reportes ejecutivos y exportaciones confiables."],
        ],
    )

    heading("13. Criterios De Aceptación")
    table(
        ["Área", "Criterio"],
        [
            ["Estabilidad", "Las rutas principales no deben mostrar pantalla 500; todo error externo debe quedar aislado con Reintentar."],
            ["Permisos", "Alumno, docente y administrador ven solo acciones autorizadas; no hay controles administrativos en HTML del alumno."],
            ["Evaluaciones", "Sesiones sin actividad no muestran pestaña Evaluaciones; hitos muestran plazo, peso, estado y nota cuando corresponda."],
            ["Encuestas", "Después de responder, el alumno ve Respuesta enviada y no se le vuelve a pedir completar la misma encuesta."],
            ["Videos", "MP4 y TXT quedan asociados a la sesión; alumno ve grabación y Chat de la clase si existe."],
            ["Certificados", "Alumno ve su certificado dentro del curso sin datos de terceros; admin puede sincronizar y enviar."],
            ["Asistencia", "Reglas diferenciadas por rol y conciliación Zoom auditada antes de uso oficial."],
            ["Rendimiento", "Vistas cacheadas menores a 500 ms; primera carga fría menor a 2 s en ambiente equivalente a producción."],
            ["Diseño", "Tres tarjetas por fila en escritorio, textos claros, contraste AA, controles de 44 px y consistencia visual."],
        ],
    )

    heading("14. Recomendación Final")
    callout(
        "Decisión sugerida",
        "El proyecto está suficientemente avanzado para preparar un piloto controlado. Antes de producción abierta, conviene cerrar estabilización, deploy VPS, automatización básica de logs/monitoreo y una batería QA por rol. La inversión más rentable después del piloto será automatizar videos de Zoom y consolidar reportes de asistencia, encuestas y certificados.",
        GREEN,
    )

    para("Documento generado como actualización del Manual de Funcionalidades Aula Virtual, usando la plantilla indicada por el usuario y el estado actual del proyecto local.", color=MUTED, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    for section in doc.sections:
        for child in list(section.header._element):
            section.header._element.remove(child)
        for child in list(section.footer._element):
            section.footer._element.remove(child)

        hp = section.header.add_paragraph()
        hp.text = "Smart Data - Aula Virtual"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.name = "Georgia"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)
        fp = section.footer.add_paragraph()
        fp.text = "Manual actualizado - uso interno"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.name = "Georgia"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)

    core = doc.core_properties
    core.title = "Manual de Funcionalidades Aula Virtual Smart Data"
    core.subject = "Estado actual, flujos operativos y plan de implementación"
    core.author = "Smart Data / Codex"
    core.comments = "Actualizado desde plantilla artifact-template-experiment-analysis."
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
